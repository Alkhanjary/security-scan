"""Turns a saved scan history record back into a downloadable report.

Markdown reuses scanner.py's own generate_markdown_report — reconstructing
Finding/ScanResult objects from the serialized dict gives byte-for-byte the
same report format the CLI's `--report` flag produces. The other formats are
built directly from the stored dict since there's no equivalent CLI output to
stay consistent with.

PDF/DOCX/XLSX depend on optional third-party libraries; each builder raises
ExportUnavailable when its library isn't installed so the UI can hide (or
explain) that format instead of returning a broken file.
"""
import csv
import datetime
import io
import json
from collections import defaultdict

from scanner import Finding, ScanResult, generate_markdown_report

CSV_FIELDS = [
    "severity", "category", "rule", "file", "line", "description",
    "impact", "improvement", "source", "ai_verdict", "ai_reason",
    "likely_test_fixture",
]

SEVERITIES = ["critical", "high", "medium", "low"]
SEVERITY_ORDER = {s: i for i, s in enumerate(SEVERITIES)}


class ExportUnavailable(RuntimeError):
    """Raised when a format's optional dependency isn't installed."""


def _is_dismissed(f: dict, include_test_files: bool = False) -> bool:
    """Mirrors scan_jobs._is_dismissed / scanner.py's own policy: an explicit
    AI verdict wins either way; failing that, a test/fixture-file finding is
    set aside unless the caller asked for those to count."""
    if f.get("ai_verdict") == "false_positive":
        return True
    if f.get("ai_verdict") == "true_positive":
        return False
    return bool(f.get("likely_test_fixture")) and not include_test_files


def _dismiss_reason(f: dict) -> tuple:
    """(tag, reason) — same three cases scanner.py's dismissed-findings
    printer uses, so the appendix reads the same as the CLI/Markdown report."""
    if f.get("ai_verdict") == "false_positive":
        return "AI", f.get("ai_reason") or ""
    if f.get("ai_reason"):
        return "AI*", f["ai_reason"]
    return "guess", "test/fixture path heuristic, unreviewed"


def _sorted_findings(record: dict, include_dismissed: bool = False) -> list:
    include_test_files = record.get("include_test_files", False)
    findings = record.get("findings", [])
    if not include_dismissed:
        findings = [f for f in findings if not _is_dismissed(f, include_test_files)]
    return sorted(findings, key=lambda f: (
        SEVERITY_ORDER.get((f.get("severity") or "").lower(), 9),
        f.get("file") or "",
        f.get("line") or 0,
    ))


def _worst_severity(findings: list) -> str:
    best_rank, best = 99, "low"
    for f in findings:
        sev = (f.get("severity") or "low").lower()
        rank = SEVERITY_ORDER.get(sev, 9)
        if rank < best_rank:
            best_rank, best = rank, sev
    return best


def report_model(record: dict) -> dict:
    """Every derived number/table used by the PDF, HTML and DOCX report
    builders, computed once so the three can't drift from each other."""
    include_test_files = record.get("include_test_files", False)
    all_findings = record.get("findings", [])
    main = [f for f in all_findings if not _is_dismissed(f, include_test_files)]
    dismissed = [f for f in all_findings if _is_dismissed(f, include_test_files)]
    main_sorted = sorted(main, key=lambda f: (
        SEVERITY_ORDER.get((f.get("severity") or "").lower(), 9),
        f.get("file") or "", f.get("line") or 0))

    files_scanned = record.get("files_scanned", 0)
    files_with_findings = len({f.get("file") for f in all_findings})

    sev_counts = {s: 0 for s in SEVERITIES}
    for f in main:
        sev = (f.get("severity") or "").lower()
        if sev in sev_counts:
            sev_counts[sev] += 1
    total_main = sum(sev_counts.values())

    severity_rows = [
        {"severity": s, "label": s.capitalize(), "count": sev_counts[s],
         "share": round(sev_counts[s] / total_main * 100) if total_main else 0}
        for s in SEVERITIES
    ]

    by_category = defaultdict(list)
    for f in main:
        by_category[f.get("category") or f.get("rule") or "uncategorized"].append(f)
    category_rows = sorted(
        ({"category": cat, "count": len(flist), "worst": _worst_severity(flist)}
         for cat, flist in by_category.items()),
        key=lambda r: (SEVERITY_ORDER.get(r["worst"], 9), -r["count"]))

    by_file = defaultdict(list)
    for f in main:
        by_file[f.get("file") or ""].append(f)
    file_rows = sorted(
        ({"file": fname, "count": len(flist), "worst": _worst_severity(flist)}
         for fname, flist in by_file.items()),
        key=lambda r: (SEVERITY_ORDER.get(r["worst"], 9), -r["count"]))

    ai_found = sum(1 for f in main if f.get("source") == "ai")

    dismissed_by_file = defaultdict(list)
    for f in dismissed:
        dismissed_by_file[f.get("file") or ""].append(f)
    dismissed_groups = []
    for fname, flist in dismissed_by_file.items():
        flist_sorted = sorted(flist, key=lambda f: (
            SEVERITY_ORDER.get((f.get("severity") or "").lower(), 9), f.get("line") or 0))
        rows = []
        for f in flist_sorted:
            tag, reason = _dismiss_reason(f)
            rows.append({
                "line": f.get("line") or "",
                "severity": (f.get("severity") or "low").lower(),
                "category": f.get("category") or f.get("rule") or "",
                "tag": tag,
                "reason": reason,
            })
        dismissed_groups.append({
            "file": fname, "count": len(flist), "worst": _worst_severity(flist), "rows": rows,
        })
    dismissed_groups.sort(key=lambda g: (SEVERITY_ORDER.get(g["worst"], 9), -g["count"]))

    if sev_counts["critical"]:
        verdict = "Critical risk. Critical issues are present and should be fixed before release."
    elif sev_counts["high"]:
        verdict = "High risk. High-severity issues should be fixed before release."
    elif total_main:
        verdict = "Moderate risk. Review the findings below."
    else:
        verdict = "No confirmed findings. This is not a guarantee of security — only that nothing matched, or everything that matched was dismissed."

    start_here = main_sorted[0] if main_sorted else None
    skipped_files = record.get("skipped_files") or []

    return {
        "target": record.get("target", ""),
        "name": record.get("name") or "",
        "generated": datetime.datetime.now().strftime("%Y-%m-%d"),
        "scanned_at": _scanned_at(record),
        "ai_used": record.get("ai_used", False),
        "ai_risk_summary": record.get("ai_risk_summary"),
        "ai_recommendations": record.get("ai_recommendations") or [],
        "findings": main_sorted,
        "dismissed_groups": dismissed_groups,
        "verdict": verdict,
        "start_here": start_here,
        "scope": {
            "files_scanned": files_scanned,
            "files_with_findings": files_with_findings,
            "raw_detections": len(all_findings),
            "findings_needing_review": total_main,
            "dismissed": len(dismissed),
            "ai_found": ai_found,
            "total_findings": total_main,
            "exit_code": record.get("exit_code"),
            "skipped_files": len(skipped_files),
        },
        "severity_rows": severity_rows,
        "category_rows": category_rows,
        "file_rows": file_rows,
        "sev_counts": sev_counts,
    }


def _scanned_at(record: dict) -> str:
    return datetime.datetime.fromtimestamp(
        record.get("timestamp", 0)).strftime("%Y-%m-%d %H:%M:%S")


def _findings_to_dataclasses(findings: list) -> list:
    out = []
    for f in findings:
        out.append(Finding(
            rule=f.get("rule", ""),
            severity=f.get("severity", "low"),
            file=f.get("file", ""),
            line=f.get("line") or 0,
            display_line=f.get("evidence", ""),
            description=f.get("description", ""),
            source=f.get("source", "regex"),
            likely_test_fixture=bool(f.get("likely_test_fixture", False)),
            ai_verdict=f.get("ai_verdict"),
            ai_reason=f.get("ai_reason"),
            impact=f.get("impact", ""),
            improvement=f.get("improvement", ""),
        ))
    return out


def build_markdown(record: dict) -> str:
    result = ScanResult(
        findings=_findings_to_dataclasses(record.get("findings", [])),
        files_scanned=record.get("files_scanned", 0),
    )
    return generate_markdown_report(
        result,
        target=record.get("target", ""),
        used_ai=record.get("ai_used", False),
        risk_summary=record.get("ai_risk_summary"),
        recommendations=record.get("ai_recommendations"),
        ai_error=None,
        exit_code=record.get("exit_code") or 0,
        threshold=record.get("fail_on", "high"),
    )


def build_json(record: dict) -> str:
    return json.dumps(record, indent=2)


def build_csv(record: dict) -> str:
    buf = io.StringIO()
    writer = csv.DictWriter(buf, fieldnames=CSV_FIELDS, extrasaction="ignore")
    writer.writeheader()
    for f in record.get("findings", []):
        row = dict(f)
        row["category"] = f.get("category") or f.get("rule", "")
        writer.writerow(row)
    return buf.getvalue()


def build_sarif(record: dict) -> str:
    """SARIF 2.1.0 — the format GitHub code scanning, Azure DevOps and most
    security dashboards ingest, so findings can be uploaded rather than
    read by hand."""
    findings = _sorted_findings(record)
    rules = {}
    results = []
    sarif_level = {"critical": "error", "high": "error", "medium": "warning", "low": "note"}

    for f in findings:
        rule_id = f.get("rule") or "finding"
        if rule_id not in rules:
            rules[rule_id] = {
                "id": rule_id,
                "name": rule_id,
                "shortDescription": {"text": f.get("description") or rule_id},
                "fullDescription": {"text": f.get("impact") or f.get("description") or rule_id},
                "help": {"text": f.get("improvement") or ""},
                "properties": {
                    "category": f.get("category") or rule_id,
                    "security-severity": {
                        "critical": "9.5", "high": "7.5", "medium": "5.0", "low": "3.0"
                    }.get((f.get("severity") or "").lower(), "3.0"),
                },
            }
        results.append({
            "ruleId": rule_id,
            "level": sarif_level.get((f.get("severity") or "").lower(), "note"),
            "message": {"text": f.get("description") or rule_id},
            "locations": [{
                "physicalLocation": {
                    "artifactLocation": {"uri": str(f.get("file") or "").replace("\\", "/")},
                    "region": {"startLine": max(1, int(f.get("line") or 1))},
                }
            }],
            "properties": {
                "severity": f.get("severity"),
                "source": f.get("source"),
                "ai_verdict": f.get("ai_verdict"),
                "ai_reason": f.get("ai_reason"),
            },
        })

    doc = {
        "$schema": "https://json.schemastore.org/sarif-2.1.0.json",
        "version": "2.1.0",
        "runs": [{
            "tool": {"driver": {
                "name": "security-scan",
                "informationUri": "https://github.com/",
                "rules": list(rules.values()),
            }},
            "results": results,
            "invocations": [{
                "executionSuccessful": True,
                "endTimeUtc": datetime.datetime.fromtimestamp(
                    record.get("timestamp", 0),
                    tz=datetime.timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
            }],
        }],
    }
    return json.dumps(doc, indent=2)


def build_xlsx(record: dict) -> bytes:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
    except ImportError:
        raise ExportUnavailable("XLSX export needs the 'openpyxl' package (pip install openpyxl).")

    wb = Workbook()

    summary = wb.active
    summary.title = "Summary"
    counts = {"critical": 0, "high": 0, "medium": 0, "low": 0}
    for f in _sorted_findings(record):
        sev = (f.get("severity") or "").lower()
        if sev in counts:
            counts[sev] += 1
    rows = [
        ("Target", record.get("target", "")),
        ("Scanned", _scanned_at(record)),
        ("Scan types", ", ".join(record.get("scan_types") or [])),
        ("Files scanned", record.get("files_scanned", 0)),
        ("AI verification", "on" if record.get("ai_used") else "off"),
        ("", ""),
        ("Critical", counts["critical"]),
        ("High", counts["high"]),
        ("Medium", counts["medium"]),
        ("Low", counts["low"]),
    ]
    for label, value in rows:
        summary.append([label, value])
    for cell in summary["A"]:
        cell.font = Font(bold=True)
    summary.column_dimensions["A"].width = 18
    summary.column_dimensions["B"].width = 70

    ws = wb.create_sheet("Findings")
    headers = ["Severity", "Category", "Rule", "File", "Line", "Description",
               "Impact", "Fix", "Source", "AI verdict", "AI reason"]
    ws.append(headers)
    header_fill = PatternFill("solid", start_color="1F2937")
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = header_fill

    sev_fill = {
        "critical": PatternFill("solid", start_color="FEE2E2"),
        "high": PatternFill("solid", start_color="FFEDD5"),
        "medium": PatternFill("solid", start_color="FEF3C7"),
        "low": PatternFill("solid", start_color="DBEAFE"),
    }
    for f in _sorted_findings(record, include_dismissed=True):
        ws.append([
            f.get("severity", ""), f.get("category") or f.get("rule", ""), f.get("rule", ""),
            f.get("file", ""), f.get("line") or "", f.get("description", ""),
            f.get("impact", ""), f.get("improvement", ""), f.get("source", ""),
            f.get("ai_verdict") or "", f.get("ai_reason") or "",
        ])
        fill = sev_fill.get((f.get("severity") or "").lower())
        if fill:
            ws.cell(row=ws.max_row, column=1).fill = fill

    for col, width in zip("ABCDEFGHIJK", (10, 22, 22, 46, 7, 46, 52, 52, 10, 15, 46)):
        ws.column_dimensions[col].width = width
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


_DOCX_SEV = {
    "critical": "DC2626", "high": "EA580C", "medium": "D97706", "low": "16A34A",
}
_DOCX_SEV_BAND = {
    "critical": "7F1D1D", "high": "9A3412", "medium": "92400E", "low": "14532D",
}
_DOCX_SEV_TINT = {
    "critical": "FEF2F2", "high": "FFF7ED", "medium": "FFFBEB", "low": "F0FDF4",
}
_DOCX_NAVY = "0F172A"


def _docx_shade(cell, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _docx_run(paragraph, text, bold=False, color=None, size=None, mono=False):
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        from docx.shared import Pt
        run.font.size = Pt(size)
    if mono:
        run.font.name = "Consolas"
    return run


def build_docx(record: dict) -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ExportUnavailable("Word export needs the 'python-docx' package (pip install python-docx).")

    model = report_model(record)
    doc = Document()

    # ------------------------------------------------------------- cover
    cover = doc.add_table(rows=1, cols=1)
    cell = cover.cell(0, 0)
    _docx_shade(cell, _DOCX_NAVY)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _docx_run(p, "SECURITY · CONFIDENTIAL", color="94A3B8", size=9)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _docx_run(p2, "VULNERABILITY ASSESSMENT REPORT", bold=True, color="FFFFFF", size=20)
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = model["scope"]
    _docx_run(p3, f"{model['target']} · {s['files_scanned']} files scanned · "
                  f"{s['findings_needing_review']} findings to review ({s['dismissed']} dismissed) · "
                  f"generated {model['generated']}", color="CBD5E1", size=9)

    stats_table = doc.add_table(rows=2, cols=6)
    stats = [
        ("FILES SCANNED", s["files_scanned"]),
        ("TOTAL FINDINGS", s["findings_needing_review"]),
        ("CRITICAL", model["sev_counts"]["critical"]),
        ("HIGH", model["sev_counts"]["high"]),
        ("MEDIUM", model["sev_counts"]["medium"]),
        ("LOW", model["sev_counts"]["low"]),
    ]
    for i, (label, value) in enumerate(stats):
        num_p = stats_table.cell(0, i).paragraphs[0]
        num_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _docx_run(num_p, str(value), bold=True, size=18,
                  color=_DOCX_SEV.get(label.lower(), "0F172A") if label in
                  ("CRITICAL", "HIGH", "MEDIUM", "LOW") else "334155")
        lab_p = stats_table.cell(1, i).paragraphs[0]
        lab_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _docx_run(lab_p, label, color="64748B", size=7)
        _docx_shade(stats_table.cell(0, i), "F1F5F9")
        _docx_shade(stats_table.cell(1, i), "F1F5F9")

    doc.add_paragraph()

    # ------------------------------------------------------ findings summary
    doc.add_heading("Findings Summary", level=1)
    doc.add_paragraph(f"Verdict: {model['verdict']}")
    if model["start_here"]:
        sf = model["start_here"]
        loc = str(sf.get("file") or "") + (f":{sf['line']}" if sf.get("line") else "")
        sh = doc.add_paragraph()
        _docx_run(sh, "Start here — ", bold=True)
        _docx_run(sh, f"{loc} — {(sf.get('severity') or '').capitalize()}: "
                      f"{sf.get('description') or sf.get('rule') or ''}")

    doc.add_heading("Scope", level=2)
    scope_tbl = doc.add_table(rows=1, cols=2)
    scope_tbl.style = "Light Grid Accent 1"
    scope_tbl.rows[0].cells[0].text, scope_tbl.rows[0].cells[1].text = "Metric", "Value"
    for label, value in [
        ("Files scanned", s["files_scanned"]),
        ("Files with findings", s["files_with_findings"]),
        ("Raw detections", s["raw_detections"]),
        ("Findings needing review", s["findings_needing_review"]),
        ("Dismissed as likely false positives", f"{s['dismissed']} (see appendix)" if s["dismissed"] else 0),
        ("Found by AI review", f"{s['ai_found']} of {s['raw_detections']}" if model["ai_used"] else "AI off"),
        ("Scanner exit code", s["exit_code"] if s["exit_code"] is not None else "—"),
    ]:
        row = scope_tbl.add_row().cells
        row[0].text, row[1].text = str(label), str(value)

    doc.add_heading("By severity", level=2)
    sev_tbl = doc.add_table(rows=1, cols=3)
    sev_tbl.style = "Light Grid Accent 1"
    for i, h in enumerate(["Severity", "Count", "Share"]):
        sev_tbl.rows[0].cells[i].text = h
    for row_data in model["severity_rows"]:
        row = sev_tbl.add_row().cells
        p = row[0].paragraphs[0]
        _docx_run(p, row_data["label"], color=_DOCX_SEV[row_data["severity"]], bold=True)
        row[1].text = str(row_data["count"])
        row[2].text = f"{row_data['share']}%"

    if model["category_rows"]:
        doc.add_heading("By category", level=2)
        cat_tbl = doc.add_table(rows=1, cols=3)
        cat_tbl.style = "Light Grid Accent 1"
        for i, h in enumerate(["Category", "Findings", "Worst severity"]):
            cat_tbl.rows[0].cells[i].text = h
        for row_data in model["category_rows"]:
            row = cat_tbl.add_row().cells
            row[0].text = row_data["category"]
            row[1].text = str(row_data["count"])
            _docx_run(row[2].paragraphs[0], row_data["worst"].capitalize(),
                      color=_DOCX_SEV[row_data["worst"]], bold=True)

    if model["file_rows"]:
        doc.add_heading("Most affected files", level=2)
        file_tbl = doc.add_table(rows=1, cols=3)
        file_tbl.style = "Light Grid Accent 1"
        for i, h in enumerate(["File", "Findings", "Worst severity"]):
            file_tbl.rows[0].cells[i].text = h
        for row_data in model["file_rows"][:15]:
            row = file_tbl.add_row().cells
            row[0].text = row_data["file"]
            row[1].text = str(row_data["count"])
            _docx_run(row[2].paragraphs[0], row_data["worst"].capitalize(),
                      color=_DOCX_SEV[row_data["worst"]], bold=True)

    if s["skipped_files"]:
        doc.add_heading("Scan notes", level=2)
        doc.add_paragraph(f"{s['skipped_files']} file(s) were skipped during the scan "
                           "(unsupported type, too large, or unreadable).", style="List Bullet")

    if model["ai_used"] and model["ai_risk_summary"]:
        doc.add_heading("AI risk summary", level=2)
        doc.add_paragraph(model["ai_risk_summary"])
        if model["ai_recommendations"]:
            for rec in model["ai_recommendations"]:
                doc.add_paragraph(str(rec), style="List Bullet")

    # ---------------------------------------------------- detailed findings
    doc.add_page_break()
    doc.add_heading("Detailed Findings", level=1)
    if not model["findings"]:
        doc.add_paragraph("No confirmed findings.")

    by_sev = defaultdict(list)
    for f in model["findings"]:
        by_sev[(f.get("severity") or "low").lower()].append(f)

    n_counter = 0
    for sev in SEVERITIES:
        flist = by_sev.get(sev)
        if not flist:
            continue
        band = doc.add_table(rows=1, cols=1)
        bcell = band.cell(0, 0)
        _docx_shade(bcell, _DOCX_SEV_BAND[sev])
        _docx_run(bcell.paragraphs[0], f"{sev.capitalize()} severity ({len(flist)})",
                  bold=True, color="FFFFFF")

        for f in flist:
            n_counter += 1
            loc = str(f.get("file") or "") + (f":{f['line']}" if f.get("line") else "")

            head_tbl = doc.add_table(rows=1, cols=1)
            hcell = head_tbl.cell(0, 0)
            _docx_shade(hcell, _DOCX_NAVY)
            _docx_run(hcell.paragraphs[0], f"#{n_counter} · {loc} — {f.get('rule') or ''}",
                      bold=True, color="FFFFFF")

            body_tbl = doc.add_table(rows=1, cols=1)
            bodycell = body_tbl.cell(0, 0)
            _docx_shade(bodycell, _DOCX_SEV_TINT[sev])
            bp = bodycell.paragraphs[0]
            _docx_run(bp, "Location: ", bold=True)
            _docx_run(bp, loc)
            p = bodycell.add_paragraph()
            _docx_run(p, "Issue: ", bold=True)
            _docx_run(p, str(f.get("description") or ""))
            if f.get("evidence"):
                p = bodycell.add_paragraph()
                _docx_run(p, "Evidence:", bold=True)
                p = bodycell.add_paragraph()
                _docx_run(p, str(f["evidence"])[:400], mono=True, size=8.5)
            if f.get("impact"):
                p = bodycell.add_paragraph()
                _docx_run(p, "Impact: ", bold=True)
                _docx_run(p, str(f["impact"]))
            if f.get("improvement"):
                p = bodycell.add_paragraph()
                _docx_run(p, "Fix: ", bold=True)
                _docx_run(p, str(f["improvement"]))
            if f.get("ai_verdict"):
                p = bodycell.add_paragraph()
                _docx_run(p, "AI verdict: ", bold=True, color="64748B")
                _docx_run(p, str(f["ai_verdict"]) + (f" ({f['ai_reason']})" if f.get("ai_reason") else ""),
                          color="64748B")
            doc.add_paragraph()

    # ------------------------------------------------ appendix (dismissed)
    if model["dismissed_groups"]:
        doc.add_page_break()
        doc.add_heading("Appendix — Possible False Positives", level=1)
        doc.add_paragraph(
            f"{s['dismissed']} detection(s) were dismissed by AI review or matched inside test "
            "fixtures, and are excluded from the counts above. They are listed here rather than "
            "dropped: this is a best guess, not proof, so scan the list before trusting it.")
        for group in model["dismissed_groups"]:
            doc.add_heading(f"{group['file']} ({group['count']} dismissed)", level=2)
            tbl = doc.add_table(rows=1, cols=4)
            tbl.style = "Light Grid Accent 1"
            for i, h in enumerate(["Line", "Severity", "Category", "Why it was dismissed"]):
                tbl.rows[0].cells[i].text = h
            for r in group["rows"]:
                row = tbl.add_row().cells
                row[0].text = str(r["line"])
                _docx_run(row[1].paragraphs[0], r["severity"].capitalize(),
                          color=_DOCX_SEV[r["severity"]], bold=True)
                row[2].text = r["category"]
                row[3].text = r["reason"][:220]

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# Severity palette shared by the cover, section bands, badges and card tints.
_PDF_SEV = {
    "critical": {"band": "#7f1d1d", "chip": "#dc2626", "tint": "#fef2f2", "text": "#dc2626"},
    "high":     {"band": "#9a3412", "chip": "#ea580c", "tint": "#fff7ed", "text": "#ea580c"},
    "medium":   {"band": "#92400e", "chip": "#d97706", "tint": "#fffbeb", "text": "#d97706"},
    "low":      {"band": "#14532d", "chip": "#16a34a", "tint": "#f0fdf4", "text": "#16a34a"},
}
_PDF_NAVY = "#0f172a"
_PDF_ACCENT = "#3b82f6"


def build_pdf(record: dict) -> bytes:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                        Preformatted, KeepTogether, Table, TableStyle,
                                        PageBreak)
    except ImportError:
        raise ExportUnavailable("PDF export needs the 'reportlab' package (pip install reportlab).")

    from xml.sax.saxutils import escape

    model = report_model(record)
    page_w, page_h = A4
    margin = 18 * mm
    content_w = page_w - 2 * margin

    styles = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=styles["BodyText"], fontSize=9.5, leading=13)
    small = ParagraphStyle("small", parent=body, fontSize=8.5, textColor=colors.HexColor("#475569"))
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], spaceBefore=14, spaceAfter=8,
                         textColor=colors.HexColor("#0f172a"))
    h3 = ParagraphStyle("h3", parent=styles["Heading3"], spaceBefore=4, spaceAfter=2)
    mono = ParagraphStyle("mono", parent=styles["Code"], fontSize=8, leading=11,
                           backColor=colors.HexColor("#eef2f9"), borderPadding=4)
    table_header_style = TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef2f9")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0f172a")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8.5),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ])

    # --------------------------------------------------------- cover (p. 1)
    def draw_cover(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(colors.HexColor(_PDF_NAVY))
        canvas.rect(0, 0, page_w, page_h, stroke=0, fill=1)

        cx = page_w / 2
        canvas.setFillColor(colors.HexColor("#94a3b8"))
        canvas.setFont("Helvetica", 9)
        canvas.drawCentredString(cx, page_h - 90, "SECURITY  ·  CONFIDENTIAL")

        canvas.setFillColor(colors.white)
        canvas.setFont("Helvetica-Bold", 30)
        canvas.drawCentredString(cx, page_h - 150, "VULNERABILITY")
        canvas.drawCentredString(cx, page_h - 190, "ASSESSMENT REPORT")

        subtitle = (f"{model['target']}  ·  {model['scope']['files_scanned']} files scanned  ·  "
                    f"{model['scope']['findings_needing_review']} findings to review "
                    f"({model['scope']['dismissed']} dismissed)  ·  generated {model['generated']}")
        canvas.setFillColor(colors.HexColor("#cbd5e1"))
        canvas.setFont("Helvetica", 9.5)
        canvas.drawCentredString(cx, page_h - 240, subtitle[:130])

        canvas.setStrokeColor(colors.HexColor(_PDF_ACCENT))
        canvas.setLineWidth(2)
        canvas.line(cx - 60, page_h - 255, cx + 60, page_h - 255)

        # stat band
        stats = [
            (str(model["scope"]["files_scanned"]), "FILES SCANNED", "#94a3b8"),
            (str(model["scope"]["findings_needing_review"]), "TOTAL FINDINGS", "#ffffff"),
            (str(model["sev_counts"]["critical"]), "CRITICAL", _PDF_SEV["critical"]["chip"]),
            (str(model["sev_counts"]["high"]), "HIGH", _PDF_SEV["high"]["chip"]),
            (str(model["sev_counts"]["medium"]), "MEDIUM", _PDF_SEV["medium"]["chip"]),
            (str(model["sev_counts"]["low"]), "LOW", _PDF_SEV["low"]["chip"]),
        ]
        band_top = 200
        canvas.setFillColor(colors.HexColor("#0b1220"))
        canvas.rect(0, 0, page_w, band_top, stroke=0, fill=1)
        n = len(stats)
        col_w = page_w / n
        for i, (num, label, color) in enumerate(stats):
            col_cx = col_w * i + col_w / 2
            canvas.setFillColor(colors.HexColor(color))
            canvas.setFont("Helvetica-Bold", 22)
            canvas.drawCentredString(col_cx, band_top - 90, num)
            canvas.setFillColor(colors.HexColor("#94a3b8"))
            canvas.setFont("Helvetica", 7.5)
            canvas.drawCentredString(col_cx, band_top - 115, label)
        canvas.restoreState()

    # ---------------------------------------------- running header / footer
    def draw_chrome(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(colors.HexColor(_PDF_NAVY))
        canvas.rect(0, page_h - 26, page_w, 26, stroke=0, fill=1)
        canvas.setFillColor(colors.white)
        canvas.setFont("Helvetica-Bold", 9)
        canvas.drawString(margin, page_h - 17, "Vulnerability Assessment Report")
        canvas.setFont("Helvetica", 8)
        canvas.drawRightString(page_w - margin, page_h - 17, "CONFIDENTIAL")

        canvas.setStrokeColor(colors.HexColor("#e2e8f0"))
        canvas.line(margin, 24, page_w - margin, 24)
        canvas.setFillColor(colors.HexColor("#94a3b8"))
        canvas.setFont("Helvetica", 7.5)
        canvas.drawString(margin, 14, str(model["target"])[:90])
        canvas.drawRightString(page_w - margin, 14, f"Page {doc.page}")
        canvas.restoreState()

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=margin, rightMargin=margin, topMargin=34 * mm, bottomMargin=16 * mm,
        title="Vulnerability Assessment Report",
    )

    story = [PageBreak()]

    # ------------------------------------------------------ findings summary
    story.append(Paragraph("Findings Summary", h2))
    story.append(Paragraph(f"&gt; Verdict: {escape(model['verdict'])}", body))
    story.append(Spacer(1, 6))
    if model["start_here"]:
        f = model["start_here"]
        loc = str(f.get("file") or "") + (f":{f['line']}" if f.get("line") else "")
        story.append(Paragraph(
            f"<b>Start here</b> — {escape(loc)} — <b>{escape((f.get('severity') or '').capitalize())}:</b> "
            f"{escape(str(f.get('description') or f.get('rule') or ''))}", body))
        story.append(Spacer(1, 10))

    story.append(Paragraph("Scope", h3))
    scope_rows = [["Metric", "Value"]]
    s = model["scope"]
    scope_rows += [
        ["Files scanned", str(s["files_scanned"])],
        ["Files with findings", str(s["files_with_findings"])],
        ["Raw detections", str(s["raw_detections"])],
        ["Findings needing review", str(s["findings_needing_review"])],
        ["Dismissed as likely false positives", f"{s['dismissed']} (see appendix)" if s["dismissed"] else "0"],
        ["Found by AI review", f"{s['ai_found']} of {s['raw_detections']}" if model["ai_used"] else "AI off"],
        ["Scanner exit code", str(s["exit_code"]) if s["exit_code"] is not None else "—"],
    ]
    t = Table(scope_rows, colWidths=[content_w * 0.55, content_w * 0.45])
    t.setStyle(table_header_style)
    story.append(t)
    story.append(Spacer(1, 12))

    story.append(Paragraph("By severity", h3))
    sev_rows = [["Severity", "Count", "Share"]]
    sev_style_extra = []
    for i, row in enumerate(model["severity_rows"], start=1):
        sev_rows.append([row["label"], str(row["count"]), f"{row['share']}%"])
        sev_style_extra.append(("TEXTCOLOR", (0, i), (0, i),
                                 colors.HexColor(_PDF_SEV[row["severity"]]["text"])))
    t = Table(sev_rows, colWidths=[content_w * 0.4, content_w * 0.3, content_w * 0.3])
    t.setStyle(TableStyle(table_header_style.getCommands() + sev_style_extra))
    story.append(t)
    story.append(Spacer(1, 12))

    if model["category_rows"]:
        story.append(Paragraph("By category", h3))
        cat_rows = [["Category", "Findings", "Worst severity"]]
        cat_extra = []
        for i, row in enumerate(model["category_rows"], start=1):
            cat_rows.append([row["category"], str(row["count"]), row["worst"].capitalize()])
            cat_extra.append(("TEXTCOLOR", (2, i), (2, i), colors.HexColor(_PDF_SEV[row["worst"]]["text"])))
        t = Table(cat_rows, colWidths=[content_w * 0.5, content_w * 0.2, content_w * 0.3])
        t.setStyle(TableStyle(table_header_style.getCommands() + cat_extra))
        story.append(t)
        story.append(Spacer(1, 12))

    if model["file_rows"]:
        story.append(Paragraph("Most affected files", h3))
        file_rows = [["File", "Findings", "Worst severity"]]
        file_extra = []
        for i, row in enumerate(model["file_rows"][:15], start=1):
            file_rows.append([row["file"], str(row["count"]), row["worst"].capitalize()])
            file_extra.append(("TEXTCOLOR", (2, i), (2, i), colors.HexColor(_PDF_SEV[row["worst"]]["text"])))
        t = Table(file_rows, colWidths=[content_w * 0.5, content_w * 0.2, content_w * 0.3])
        t.setStyle(TableStyle(table_header_style.getCommands() + file_extra))
        story.append(t)
        story.append(Spacer(1, 12))

    if s["skipped_files"]:
        story.append(Paragraph("Scan notes", h3))
        story.append(Paragraph(
            f"• {s['skipped_files']} file(s) were skipped during the scan "
            "(unsupported type, too large, or unreadable).", body))
        story.append(Spacer(1, 8))

    if model["ai_used"] and model["ai_risk_summary"]:
        story.append(Paragraph("AI Risk Summary", h3))
        story.append(Paragraph(escape(str(model["ai_risk_summary"])), body))
        if model["ai_recommendations"]:
            for rec in model["ai_recommendations"]:
                story.append(Paragraph("• " + escape(str(rec)), body))
        story.append(Spacer(1, 8))

    # ---------------------------------------------------- detailed findings
    story.append(PageBreak())
    story.append(Paragraph("Detailed Findings", h2))

    if not model["findings"]:
        story.append(Paragraph("No confirmed findings.", body))

    by_sev = defaultdict(list)
    for f in model["findings"]:
        by_sev[(f.get("severity") or "low").lower()].append(f)

    n_counter = 0
    for sev in SEVERITIES:
        flist = by_sev.get(sev)
        if not flist:
            continue
        pal = _PDF_SEV[sev]
        band = Table([[Paragraph(
            f'<font color="white"><b>{sev.capitalize()} severity ({len(flist)})</b></font>', body)]],
            colWidths=[content_w])
        band.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(pal["band"])),
            ("TOPPADDING", (0, 0), (-1, -1), 8), ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ]))
        story.append(Spacer(1, 10))
        story.append(band)
        story.append(Spacer(1, 6))

        for f in flist:
            n_counter += 1
            loc = str(f.get("file") or "") + (f":{f['line']}" if f.get("line") else "")
            title_row = Table([[Paragraph(
                f'<font color="white"><b>#{n_counter} · {escape(loc)} — {escape(f.get("rule") or "")}</b></font>',
                body)]], colWidths=[content_w])
            title_row.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(_PDF_NAVY)),
                ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
            ]))

            body_flow = []
            body_flow.append(Paragraph(f"<b>Location:</b> {escape(loc)}", body))
            body_flow.append(Paragraph(
                f"<b>Issue:</b> {escape(str(f.get('description') or ''))}", body))
            if f.get("evidence"):
                body_flow.append(Paragraph("<b>Evidence:</b>", body))
                body_flow.append(Preformatted(str(f["evidence"])[:400], mono))
            if f.get("impact"):
                body_flow.append(Paragraph(f"<b>Impact:</b> {escape(str(f['impact']))}", body))
            if f.get("improvement"):
                body_flow.append(Paragraph(f"<b>Fix:</b> {escape(str(f['improvement']))}", body))
            if f.get("ai_verdict"):
                body_flow.append(Paragraph(
                    f"<b>AI verdict:</b> {escape(f['ai_verdict'])}"
                    + (f" ({escape(str(f['ai_reason']))})" if f.get("ai_reason") else ""), small))

            body_table = Table([[body_flow]], colWidths=[content_w])
            body_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(pal["tint"])),
                ("BOX", (0, 0), (-1, -1), 1, colors.HexColor(pal["chip"])),
                ("TOPPADDING", (0, 0), (-1, -1), 8), ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("LEFTPADDING", (0, 0), (-1, -1), 10), ("RIGHTPADDING", (0, 0), (-1, -1), 10),
            ]))
            story.append(KeepTogether([title_row, body_table, Spacer(1, 8)]))

    # ------------------------------------------------ appendix (dismissed)
    if model["dismissed_groups"]:
        story.append(PageBreak())
        story.append(Paragraph("Appendix — Possible False Positives", h2))
        total_dismissed = model["scope"]["dismissed"]
        story.append(Paragraph(
            f"{total_dismissed} detection(s) were dismissed by AI review or matched inside test "
            "fixtures, and are excluded from the counts above. They are listed here rather than "
            "dropped: this is a best guess, not proof, so scan the list before trusting it.", body))
        story.append(Spacer(1, 10))

        for group in model["dismissed_groups"]:
            story.append(Paragraph(f"{escape(group['file'])} ({group['count']} dismissed)", h3))
            rows = [["Line", "Severity", "Category", "Why it was dismissed"]]
            for r in group["rows"]:
                rows.append([str(r["line"]), r["severity"].capitalize(), r["category"], r["reason"][:180]])
            t = Table(rows, colWidths=[content_w * 0.08, content_w * 0.14, content_w * 0.28, content_w * 0.5])
            style_cmds = list(table_header_style.getCommands())
            for i, r in enumerate(group["rows"], start=1):
                style_cmds.append(("TEXTCOLOR", (1, i), (1, i), colors.HexColor(_PDF_SEV[r["severity"]]["text"])))
            t.setStyle(TableStyle(style_cmds))
            story.append(t)
            story.append(Spacer(1, 10))

    doc.build(story, onFirstPage=draw_cover, onLaterPages=draw_chrome)
    return buf.getvalue()
